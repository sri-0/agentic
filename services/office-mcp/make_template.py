"""Author a minimal default docxtpl (.docx) report template with python-docx.

docxtpl renders Jinja2 tags embedded in a *real* Word document, so the template
must itself be a .docx. This builds one with a title, metadata, a summary, and a
looped sections list — enough to test render_report_docx end-to-end and to serve
as a starting point designers can rebrand.

Run standalone to (re)generate:
    ./.venv/bin/python make_template.py            # writes ./templates/report.docx
"""
from __future__ import annotations

import os

from docx import Document
from docx.shared import Pt


def build_default_template(path: str) -> str:
    """Write a default docxtpl template to `path` and return it."""
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc = Document()

    doc.add_heading("{{ title }}", level=0)

    meta = doc.add_paragraph()
    meta.add_run("Author: ").bold = True
    meta.add_run("{{ author }}")
    meta.add_run("    Date: ").bold = True
    meta.add_run("{{ date }}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph("{{ summary }}")

    # Jinja2 loop over sections. docxtpl reads {% %} tags in paragraph runs.
    doc.add_paragraph("{% for s in sections %}")
    doc.add_heading("{{ s.heading }}", level=2)
    doc.add_paragraph("{{ s.body }}")
    doc.add_paragraph("{% endfor %}")

    # A default font size so the doc looks intentional.
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.save(path)
    return path


if __name__ == "__main__":
    out = os.environ.get("OFFICE_TEMPLATE_DIR", "./templates")
    dest = os.path.join(out, "report.docx")
    build_default_template(dest)
    print(f"wrote {dest}")
